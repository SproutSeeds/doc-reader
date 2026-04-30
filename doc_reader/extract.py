from __future__ import annotations

import re
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Iterator

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md", ".markdown"}


class UnsupportedFormatError(ValueError):
    pass


@dataclass(frozen=True)
class DocumentBlock:
    text: str
    page_number: int | None = None


def iter_document_blocks(path: Path, *, preserve_all: bool = False) -> Iterator[str]:
    for block in iter_document_blocks_with_meta(path, preserve_all=preserve_all):
        yield block.text


def iter_document_blocks_with_meta(path: Path, *, preserve_all: bool = False) -> Iterator[DocumentBlock]:
    suffix = path.suffix.lower()

    if suffix not in SUPPORTED_SUFFIXES:
        supported = ", ".join(sorted(SUPPORTED_SUFFIXES))
        raise UnsupportedFormatError(
            f"Unsupported file format: {suffix or '<none>'}. Supported: {supported}"
        )

    if suffix == ".pdf":
        yield from _iter_pdf_blocks(path, preserve_all=preserve_all)
    elif suffix == ".docx":
        yield from _wrap_plain_blocks(_iter_docx_blocks(path, preserve_all=preserve_all))
    elif suffix in {".md", ".markdown"}:
        yield from _wrap_plain_blocks(_iter_markdown_blocks(path, preserve_all=preserve_all))
    else:
        yield from _wrap_plain_blocks(_iter_plaintext_blocks(path, preserve_all=preserve_all))


def _wrap_plain_blocks(blocks: Iterable[str]) -> Iterator[DocumentBlock]:
    for text in blocks:
        if text:
            yield DocumentBlock(text=text, page_number=None)


def _iter_pdf_blocks(path: Path, *, preserve_all: bool) -> Iterator[DocumentBlock]:
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "PDF support requires pypdf. Install dependencies from requirements.txt."
        ) from exc

    reader = PdfReader(str(path))
    seen_top_lines: Counter[str] = Counter()
    seen_bottom_lines: Counter[str] = Counter()

    for page_number, page in enumerate(reader.pages, start=1):
        raw_text = page.extract_text() or ""
        lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
        lines = [line for line in lines if not _looks_like_layout_noise_line(line)]
        if not lines:
            continue

        if not preserve_all:
            top = lines[0]
            if _looks_like_running_edge(top) and seen_top_lines[top] >= 2:
                lines = lines[1:]
            seen_top_lines[top] += 1

            if not lines:
                continue

            bottom = lines[-1]
            if _looks_like_running_edge(bottom) and seen_bottom_lines[bottom] >= 2:
                lines = lines[:-1]
            seen_bottom_lines[bottom] += 1

        text = _clean_inline(" ".join(lines))
        if text and (preserve_all or not _is_noise_block(text)):
            yield DocumentBlock(text=text, page_number=page_number)


def _iter_docx_blocks(path: Path, *, preserve_all: bool) -> Iterator[str]:
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "DOCX support requires python-docx. Install dependencies from requirements.txt."
        ) from exc

    doc = Document(str(path))

    for para in doc.paragraphs:
        text = _clean_inline(para.text)
        if text and (preserve_all or not _is_noise_block(text)):
            yield text

    # Tables often contain important data; read row-wise and speak as summaries.
    for table in doc.tables:
        for row in table.rows:
            cells = [_clean_inline(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if not cells:
                continue
            row_text = " | ".join(cells)
            if preserve_all or not _is_noise_block(row_text):
                yield f"Table row: {row_text}"


def _iter_plaintext_blocks(path: Path, *, preserve_all: bool) -> Iterator[str]:
    yield from _iter_lines_as_paragraphs(path, markdown_mode=False, preserve_all=preserve_all)


def _iter_markdown_blocks(path: Path, *, preserve_all: bool) -> Iterator[str]:
    yield from _iter_lines_as_paragraphs(path, markdown_mode=True, preserve_all=preserve_all)


def _iter_lines_as_paragraphs(path: Path, markdown_mode: bool, *, preserve_all: bool) -> Iterator[str]:
    buffer: list[str] = []
    in_code_block = False

    for raw_line in _iter_text_lines(path):
        line = raw_line.rstrip("\n")

        if markdown_mode and line.strip().startswith("```"):
            in_code_block = not in_code_block
            continue
        if markdown_mode and in_code_block:
            continue

        cleaned = _clean_markdown_line(line) if markdown_mode else _clean_inline(line)
        if not cleaned:
            if buffer:
                block = _clean_inline(" ".join(buffer))
                if block and (preserve_all or not _is_noise_block(block)):
                    yield block
                buffer.clear()
            continue
        buffer.append(cleaned)

    if buffer:
        block = _clean_inline(" ".join(buffer))
        if block and (preserve_all or not _is_noise_block(block)):
            yield block


def _clean_markdown_line(line: str) -> str:
    line = line.rstrip()
    if not line:
        return ""

    line = re.sub(r"^#{1,6}\s+", "", line)
    line = re.sub(r"!\[[^\]]*\]\([^)]*\)", "", line)
    line = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", line)
    line = re.sub(r"^\s*>\s?", "", line)
    line = re.sub(r"^\s*[-*+]\s+", "Bullet: ", line)
    line = re.sub(r"^\s*\d+\.\s+", "Item: ", line)
    line = line.replace("`", "")
    return _clean_inline(line)


def _clean_inline(text: str) -> str:
    if not text:
        return ""

    text = text.replace("\u00ad", "")
    text = text.replace("\xa0", " ")
    text = re.sub(r"([A-Za-z])-\s+([A-Za-z])", r"\1\2", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _is_noise_block(text: str) -> bool:
    if not text:
        return True
    if re.fullmatch(r"[\W_]+", text):
        return True
    if re.fullmatch(r"\d+", text):
        return True
    if re.fullmatch(r"(?:page\s+)?\d+(?:\s*/\s*\d+)?", text, flags=re.IGNORECASE):
        return True
    if len(text) < 2:
        return True
    return False


def _looks_like_running_edge(line: str) -> bool:
    line = line.strip()
    if not line:
        return False
    if len(line) > 90:
        return False
    if re.search(r"(?:page\s+)?\d+(?:\s*/\s*\d+)?", line, flags=re.IGNORECASE):
        return True
    return len(line.split()) <= 10


def _looks_like_layout_noise_line(line: str) -> bool:
    cleaned = _clean_inline(line)
    if not cleaned:
        return True
    if _is_noise_block(cleaned):
        return True
    if re.fullmatch(r"[•·▪◦*]+", cleaned):
        return True
    if re.fullmatch(r"[._\-—=~•·▪◦*\s]{6,}", cleaned):
        return True
    if re.fullmatch(
        r"(?:[•·▪◦*]\s*)?\d+(?:\s+\d+){2,}(?:\s*[•·▪◦*]\s*\d+(?:\s+\d+)*)*",
        cleaned,
    ):
        return True
    return False


def _iter_text_lines(path: Path) -> Iterator[str]:
    encodings = ("utf-8", "utf-8-sig", "latin-1")
    for encoding in encodings:
        try:
            with path.open("r", encoding=encoding) as handle:
                for line in handle:
                    yield line
            return
        except UnicodeDecodeError:
            continue

    with path.open("r", encoding="utf-8", errors="replace") as handle:
        for line in handle:
            yield line
